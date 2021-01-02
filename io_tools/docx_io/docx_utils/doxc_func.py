from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn


def SetStyle(title_run, size, color):
    title_run.font.size = Pt(size)  # 设置标题字体大小
    title_run.font.name = 'Times New Roman'  # 设置标题西文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置标题中文字体
    title_run.font.color.rgb = RGBColor(color[0], color[1], color[2])  # 字体颜色


def AddHeadText(document, text, size, color=None):
    if color is None:
        color = [0, 0, 0]
    title_ = document.add_heading(level=3)
    title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    SetStyle(title_run, size, color)


def AddParagraphText(document, text, size, color=None):
    if color is None:
        color = [0, 0, 0]
    title_ = document.add_paragraph()
    title_run = title_.add_run(text)  # 添加标题内容
    SetStyle(title_run, size, color)
